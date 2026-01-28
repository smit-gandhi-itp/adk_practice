from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


# =========================
# Phase 1 – CLI Helpers
# =========================

def ask_text(question: str) -> str:
    print(f"\n{question}")
    return input("> ").strip()


def ask_single_choice(question: str, options: list[str]) -> str:
    print(f"\n{question}")
    for i, opt in enumerate(options, start=1):
        print(f"{i}. {opt}")

    while True:
        choice = input("> ").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(options):
            return options[int(choice) - 1]
        print("❌ Invalid choice. Please enter a valid number.")


def ask_multi_choice(question: str, options: list[str]) -> list[str]:
    print(f"\n{question}")
    for i, opt in enumerate(options, start=1):
        print(f"{i}. {opt}")

    print("Enter comma-separated numbers (e.g. 1,3,5). Press Enter to skip.")

    while True:
        raw = input("> ").strip()
        if not raw:
            return []

        try:
            indexes = [int(i.strip()) for i in raw.split(",")]
            if all(1 <= i <= len(options) for i in indexes):
                return [options[i - 1] for i in indexes]
        except ValueError:
            pass

        print("❌ Invalid input. Try again.")


# =========================
# Phase 1 – Design Engine
# =========================

def run_phase_1_design_engine() -> dict:
    print("\n======================================")
    print(" Agentic Design Engine – Phase 1 ")
    print(" User Input & Context Setup ")
    print("======================================")

    project_name = ask_text("Project Name (e.g. Payment Gateway Platform)")

    project_type = ask_single_choice(
        "Project Type",
        [
            "Web Application",
            "Mobile Application",
            "Backend Service",
            "Data Platform",
            "ML / AI System",
            "Internal Tool",
        ],
    )

    platform = ask_single_choice(
        "Primary Platform",
        [
            "Web",
            "Mobile",
            "Backend / API",
            "Data / Analytics",
            "Mixed",
        ],
    )

    description = ask_text(
        "Project Description\nBriefly describe what this system does and why it exists"
    )

    core_features_raw = ask_text(
        "Core Features\nEnter comma-separated features (e.g. Auth, Payments, Reports)"
    )
    core_features = [f.strip() for f in core_features_raw.split(",") if f.strip()]

    expected_user_scale = ask_single_choice(
        "Expected User Scale",
        [
            "Prototype / Internal",
            "Up to 10k users",
            "10k - 100k users",
            "100k - 1M users",
            "1M+ users",
        ],
    )

    constraints = ask_multi_choice(
        "Primary Constraints",
        [
            "Performance",
            "Cost",
            "Security",
            "Compliance",
            "Time-to-market",
            "Scalability",
        ],
    )

    return {
        "project_name": project_name,
        "project_type": project_type,
        "platform": platform,
        "description": description,
        "core_features": core_features,
        "expected_user_scale": expected_user_scale,
        "constraints": constraints,
    }


# =========================
# Phase 2
# ========================


## FOR GROQ the below function ## 
def normalize_phase_2_questions(clarification_output: dict) -> dict:
    """
    Convert Groq-safe list schema into the dict format
    expected by the Phase 2 engine.
    """
    normalized = {}

    for q in clarification_output.get("questions", []):
        normalized[q["question_text"]] = {
            "type": q["type"],
            "options": q["options"]
        }

    return {"questions": normalized}


def run_phase_2_clarification_engine(phase_2_clarification_questions: dict) -> dict:
    print("\n======================================")
    print(" Agentic Design Engine - Phase 2 ")
    print(" Clarification Questions ")
    print("======================================")

    answers = {}

    questions = phase_2_clarification_questions.get("questions", {})

    for question_text, spec in questions.items():
        q_type = spec["type"]

        if q_type != "multi_choice":
            raise ValueError(
                f"Unsupported question type: {q_type}. "
                "Only multi_choice is allowed in this Phase 2 engine."
            )

        options = spec["options"]

        # Ask user to select one or more options
        selected = ask_multi_choice(question_text, options)

        # Robust check for "Other" selection
        if any("other" in choice.lower() for choice in selected):
            other_answer = ask_text(
                f"You selected 'Other' for the question:\n{question_text}\nPlease specify:"
            )
            # Replace any "Other" choice with the custom input
            selected = [
                other_answer if "other" in choice.lower() else choice for choice in selected
            ]

        answers[question_text] = selected

    return answers





# =========================
# Phase 3
# =========================



# def render_phase3_design_to_word(system_design: dict, output_path: str):
#     doc = Document()

#     # =========================
#     # Helpers
#     # =========================
#     def heading(text, level=1):
#         h = doc.add_heading(text, level)
#         h.alignment = WD_ALIGN_PARAGRAPH.LEFT

#     def para(text, bold=False):
#         p = doc.add_paragraph()
#         r = p.add_run(text)
#         r.bold = bold
#         r.font.size = Pt(11)

#     def bullets(items, fallback=None):
#         if items:
#             for i in items:
#                 if i:
#                     doc.add_paragraph(str(i), style="List Bullet")
#         elif fallback:
#             para(fallback)

#     def table(headers, rows):
#         t = doc.add_table(rows=1, cols=len(headers))
#         t.style = "Table Grid"
#         t.alignment = WD_TABLE_ALIGNMENT.CENTER
#         for i, h in enumerate(headers):
#             t.rows[0].cells[i].text = h
#         for r in rows:
#             cells = t.add_row().cells
#             for i, v in enumerate(r):
#                 cells[i].text = str(v)

#     # =========================
#     # Title Page (INTENTIONAL)
#     # =========================
#     title = system_design["executive_summary"]["title"]

#     doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
#     doc.add_paragraph("System Design Document").alignment = WD_ALIGN_PARAGRAPH.CENTER
#     doc.add_paragraph("Generated by Agentic Design Engine").alignment = WD_ALIGN_PARAGRAPH.CENTER
#     doc.add_paragraph("Version 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER

#     doc.add_page_break()

#     # =========================
#     # 1. Executive Summary
#     # =========================
#     heading("1. Executive Summary")

#     es = system_design["executive_summary"]
#     para(f"Project Type: {es['project_type']}", bold=True)
#     para(es["purpose"])

#     heading("Core Features", 2)
#     bullets(es["core_features"], "No core features explicitly specified.")

#     heading("Constraints", 2)
#     bullets(es["constraints"], "No major constraints were identified.")

#     us = es["user_scale"]
#     para(
#         f"User Scale: {us['expected_active_users']} | "
#         f"Concurrent Users: {us['concurrent_users']}"
#     )

#     # =========================
#     # 2. System Overview
#     # =========================
#     heading("2. System Overview")

#     so = system_design["system_overview"]

#     heading("Functional Goals", 2)
#     bullets(so["functional_goals"], "No explicit functional goals defined.")

#     heading("Non-Functional Requirements", 2)
#     bullets(
#         so["non_functional_requirements"],
#         "Non-functional requirements are inferred implicitly."
#     )

#     heading("Primary User Personas", 2)
#     bullets(so["primary_user_personas"], "Primary users are internal stakeholders.")

#     heading("User Flows", 2)
#     if so["user_flows"]:
#         for flow in so["user_flows"]:
#             para(flow["name"], bold=True)
#             para(flow["description"])
#             para(f"Entry Points: {', '.join(flow['entry_points'])}")
#             para(f"Success Criteria: {flow['success_criteria']}")
#     else:
#         para(
#             "No explicit user flows were defined. "
#             "System interactions are primarily automated or internally driven."
#         )

#     # =========================
#     # 3. Architecture Design
#     # =========================
#     heading("3. Architecture Design")

#     arch = system_design["architecture_design"]
#     para(arch["overview"])

#     heading("Components", 2)
#     for c in arch["components"]:
#         para(c["name"], bold=True)
#         para(f"Responsibility: {c['responsibility']}")
#         bullets(c["technologies"], "Technology stack to be finalized.")

#     heading("Design Rationale", 2)
#     bullets(
#         arch["rationale"],
#         "Architecture decisions were made to balance scalability, security, and maintainability."
#     )

#     # =========================
#     # 4. Database Design
#     # =========================
#     heading("4. Database Design")

#     db = system_design["database_design"]
#     para(f"Database Type: {db['db_type']}")
#     para(db["storage_characteristics"])

#     for s in db["schemas"]:
#         heading(f"Table: {s['table_name']}", 3)
#         para(s["description"])

#         if s["columns"]:
#             table(
#                 ["Column", "Type", "Nullable", "Description"],
#                 [
#                     [c["name"], c["type"], c["nullable"], c["description"]]
#                     for c in s["columns"]
#                 ],
#             )
#         else:
#             para("No detailed column schema provided.")

#     # =========================
#     # 5. Security & Compliance
#     # =========================
#     heading("5. Security and Compliance")

#     sec = system_design["security_and_compliance"]
#     para(f"Authentication Method: {sec['authentication']['method']}")
#     bullets(sec["authorization"], "Authorization policies to be finalized.")
#     bullets(sec["encryption"], "Encryption standards applied where applicable.")
#     bullets(sec["compliance"], "Compliance requirements depend on deployment geography.")

#     # =========================
#     # 6. Deployment Strategy
#     # =========================
#     heading("6. Deployment Strategy")

#     dep = system_design["deployment_strategy"]
#     para(dep["model"])
#     bullets(dep["containerization"], "Containerization approach under evaluation.")
#     bullets(dep["ci_cd"], "CI/CD pipelines will be established.")
#     bullets(dep["rollout_strategy"], "Gradual rollout strategy planned.")

#     # =========================
#     # 7. Scalability & Reliability
#     # =========================
#     heading("7. Scalability and Reliability")

#     sr = system_design["scalability_and_reliability"]
#     bullets(sr["load_balancing"], "Load balancing will be handled via managed services.")
#     bullets(sr["autoscaling"], "Autoscaling rules will be defined post-traffic analysis.")
#     bullets(sr["caching_strategy"], "Caching strategy will be added where required.")
#     bullets(sr["failover_and_dr"], "Disaster recovery strategy under consideration.")

#     # =========================
#     # 8. Cost Estimation
#     # =========================
#     heading("8. Cost and Resource Estimation")

#     cost = system_design["cost_and_resource_estimation"]
#     if cost["cost_items"]:
#         for c in cost["cost_items"]:
#             para(
#                 f"{c['name']}: ${c['monthly_estimate_usd']} / month – {c['rationale']}"
#             )
#     else:
#         para(
#             "Detailed cost estimation is deferred due to early-stage design assumptions."
#         )

#     # =========================
#     # 9. Testing & QA
#     # =========================
#     heading("9. Testing and QA Strategy")

#     qa = system_design["testing_and_qa_strategy"]
#     for k, v in qa.items():
#         if k == "acceptance_criteria":
#             continue
#         para(v["name"], bold=True)
#         para(v["scope"])
#         bullets(v["tools"], "Tooling to be finalized.")

#     heading("Acceptance Criteria", 2)
#     bullets(
#         qa["acceptance_criteria"],
#         "System must meet functional and non-functional requirements."
#     )

#     # =========================
#     # 10. Appendices
#     # =========================
#     heading("10. Appendices")

#     refs = system_design["appendices"]["references"]
#     if refs:
#         bullets(refs)
#     else:
#         para("No external references provided for this version of the document.")

#     doc.save(output_path)


## FOR GROQ ##

def safe_get(d: dict, path: list, default=None):
    """
    Safely access nested dictionaries.
    Example: safe_get(phase3, ["executive_summary", "title"])
    """
    cur = d
    for key in path:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)


def sanitize_mermaid_output(text: str) -> str | None:
    if not text or not isinstance(text, str):
        return None

    lowered = text.lower()

    refusal_markers = [
        "i can't help",
        "i cannot help",
        "sorry",
        "unable to",
        "as an ai",
        "cannot generate"
    ]
    if any(marker in lowered for marker in refusal_markers):
        return None

    # Remove markdown fences
    text = text.strip()
    text = re.sub(r"^```[a-zA-Z]*", "", text)
    text = re.sub(r"```$", "", text)

    lines = text.splitlines()
    cleaned_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Normalize labeled arrows
        line = re.sub(
            r"(.*?)-->\s*\|\s*(.*?)\s*\|\s*(.*)",
            r"\1-->| \2 | \3",
            line
        )

        # Normalize arrow variants
        line = re.sub(r"-{3,}>", "-->", line)
        line = re.sub(r"<-{3,}", "<--", line)
        line = re.sub(r"==>", "-->", line)

        cleaned_lines.append(line)

    if not cleaned_lines:
        return None

    first = cleaned_lines[0]
    valid_headers = (
        "flowchart", "graph", "erDiagram",
        "sequenceDiagram", "stateDiagram", "classDiagram"
    )

    if not first.startswith(valid_headers):
        cleaned_lines.insert(0, "flowchart LR")

    return "\n".join(cleaned_lines)


def render_phase3_design_to_word(phase3: dict, output_path: str):
    """
    Renders Phase 3 system design to a Word document.
    diagrams: optional dict containing Mermaid code strings:
      - diagrams['system_architecture']
      - diagrams['user_flows']
      - diagrams['database_er']
    """
    doc = Document()

    def add_heading(text, level=1):
        doc.add_heading(text, level=level)

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold

    # -----------------------------
    # Title
    # -----------------------------
    title = safe_get(phase3, ["executive_summary", "title"], "System Design Document")
    add_heading(title, 0)

    # -----------------------------
    # Executive Summary
    # -----------------------------
    es = safe_get(phase3, ["executive_summary"], {})
    if es:
        add_heading("Executive Summary", 1)
        add_paragraph(f"Project Type: {es.get('project_type', '')}")
        add_paragraph(es.get("purpose", ""), bold=True)
        add_paragraph("Core Features", bold=True)
        for f in es.get("core_features", []):
            add_paragraph(f"- {f}")
        add_paragraph("Constraints", bold=True)
        for c in es.get("constraints", []):
            add_paragraph(f"- {c}")
        us = es.get("user_scale", {})
        if us:
            add_paragraph(
                f"User Scale: {us.get('expected_active_users')} MAU, "
                f"{us.get('concurrent_users')} concurrent users"
            )

    # -----------------------------
    # System Overview
    # -----------------------------
    so = safe_get(phase3, ["system_overview"], {})
    if so:
        add_heading("System Overview", 1)
        for section, items in [
            ("Functional Goals", so.get("functional_goals", [])),
            ("Non-Functional Requirements", so.get("non_functional_requirements", [])),
            ("Primary User Personas", so.get("primary_user_personas", [])),
        ]:
            add_paragraph(section, bold=True)
            for i in items:
                add_paragraph(f"- {i}")

    # -----------------------------
    # Architecture Design (TABLE)
    # -----------------------------
    ad = safe_get(phase3, ["architecture_design"], {})
    if ad:
        add_heading("Architecture Design", 1)
        add_paragraph(ad.get("overview", ""))
        components = ad.get("components", [])
        if components:
            add_paragraph("System Components", bold=True)
            rows = []
            for c in components:
                rows.append([
                    c.get("name"),
                    c.get("responsibility"),
                    ", ".join(c.get("technologies", [])),
                    ", ".join(c.get("interfaces", [])),
                ])
            add_table(doc, headers=["Component", "Responsibility", "Technologies", "Interfaces"], rows=rows)

    # -----------------------------
    # Database Design (TABLE)
    # -----------------------------
    db = safe_get(phase3, ["database_design"], {})
    if db:
        add_heading("Database Design", 1)
        add_paragraph(f"Database Type: {db.get('db_type')}")
        add_paragraph(db.get("storage_characteristics", ""))
        for table in db.get("schemas", []):
            add_paragraph(f"Table: {table.get('table_name')}", bold=True)
            rows = []
            for col in table.get("columns", []):
                rows.append([
                    col.get("name"),
                    col.get("type"),
                    col.get("nullable"),
                    col.get("description"),
                ])
            add_table(doc, headers=["Column", "Type", "Nullable", "Description"], rows=rows)

    # -----------------------------
    # Cost Estimation (TABLE)
    # -----------------------------
    ce = safe_get(phase3, ["cost_and_resource_estimation"], {})
    if ce:
        add_heading("Cost Estimation", 1)
        rows = []
        for c in ce.get("cost_items", []):
            rows.append([c.get("name"), f"${c.get('monthly_estimate_usd')}", c.get("rationale")])
        add_table(doc, headers=["Cost Item", "Monthly Cost", "Rationale"], rows=rows)

    # -----------------------------
    # Testing Strategy
    # -----------------------------
    tq = safe_get(phase3, ["testing_and_qa_strategy"], {})
    if tq:
        add_heading("Testing & QA Strategy", 1)
        for key, t in tq.items():
            if isinstance(t, dict):
                add_paragraph(f"{t.get('name')} – {t.get('scope')}")

    # -----------------------------
    # Appendices
    # -----------------------------
    ap = safe_get(phase3, ["appendices"], {})
    if ap:
        add_heading("Appendices", 1)
        glossary = ap.get("glossary", [])
        if glossary:
            add_paragraph("Glossary", bold=True)
            for g in glossary:
                add_paragraph(f"{g.get('term')}: {g.get('definition')}")
        refs = ap.get("references", [])
        if refs:
            add_paragraph("References", bold=True)
            for r in refs:
                add_paragraph(r)
        notes = ap.get("additional_notes")
        if notes:
            add_paragraph(notes)

    # -----------------------------
    # Phase 3 Diagrams (NEW)
    # -----------------------------
    md = safe_get(phase3, ["mermaid_diagrams"], {})
    if md:
        add_heading("System Architecture & Diagrams", 1)
        if md.get("system_architecture"):
            add_paragraph("System Architecture Diagram (Mermaid Code):", bold=True)
            add_paragraph(sanitize_mermaid_output(md["system_architecture"]))
        if md.get("user_flows"):
            add_paragraph("User Flow Diagram (Mermaid Code):", bold=True)
            add_paragraph(sanitize_mermaid_output(md["user_flows"]))
        if md.get("database_er"):
            add_paragraph("Database ER Diagram (Mermaid Code):", bold=True)
            add_paragraph(sanitize_mermaid_output(md["database_er"]))

    doc.save(output_path)




## old ##

# def render_phase3_design_to_word(phase3: dict, output_path: str):
#     doc = Document()

#     def add_heading(text, level=1):
#         doc.add_heading(text, level=level)

#     def add_paragraph(text, bold=False):
#         p = doc.add_paragraph()
#         run = p.add_run(text)
#         run.bold = bold

#     # -----------------------------
#     # Title
#     # -----------------------------
#     title = safe_get(phase3, ["executive_summary", "title"], "System Design Document")
#     add_heading(title, 0)

#     # =============================
#     # Executive Summary
#     # =============================
#     es = safe_get(phase3, ["executive_summary"], {})
#     if es:
#         add_heading("Executive Summary", 1)
#         add_paragraph(f"Project Type: {es.get('project_type', '')}")
#         add_paragraph(es.get("purpose", ""), bold=True)

#         add_paragraph("Core Features", bold=True)
#         for f in es.get("core_features", []):
#             add_paragraph(f"- {f}")

#         add_paragraph("Constraints", bold=True)
#         for c in es.get("constraints", []):
#             add_paragraph(f"- {c}")

#         us = es.get("user_scale", {})
#         if us:
#             add_paragraph(
#                 f"User Scale: {us.get('expected_active_users')} MAU, "
#                 f"{us.get('concurrent_users')} concurrent users"
#             )

#     # =============================
#     # System Overview
#     # =============================
#     so = safe_get(phase3, ["system_overview"], {})
#     if so:
#         add_heading("System Overview", 1)

#         for section, items in [
#             ("Functional Goals", so.get("functional_goals", [])),
#             ("Non-Functional Requirements", so.get("non_functional_requirements", [])),
#             ("Primary User Personas", so.get("primary_user_personas", [])),
#         ]:
#             add_paragraph(section, bold=True)
#             for i in items:
#                 add_paragraph(f"- {i}")

#     # =============================
#     # Architecture Design (TABLE)
#     # =============================
#     ad = safe_get(phase3, ["architecture_design"], {})
#     if ad:
#         add_heading("Architecture Design", 1)
#         add_paragraph(ad.get("overview", ""))

#         components = ad.get("components", [])
#         if components:
#             add_paragraph("System Components", bold=True)
#             rows = []
#             for c in components:
#                 rows.append([
#                     c.get("name"),
#                     c.get("responsibility"),
#                     ", ".join(c.get("technologies", [])),
#                     ", ".join(c.get("interfaces", [])),
#                 ])

#             add_table(
#                 doc,
#                 headers=["Component", "Responsibility", "Technologies", "Interfaces"],
#                 rows=rows
#             )

#     # =============================
#     # Database Design (TABLE)
#     # =============================
#     db = safe_get(phase3, ["database_design"], {})
#     if db:
#         add_heading("Database Design", 1)
#         add_paragraph(f"Database Type: {db.get('db_type')}")
#         add_paragraph(db.get("storage_characteristics", ""))

#         for table in db.get("schemas", []):
#             add_paragraph(f"Table: {table.get('table_name')}", bold=True)
#             rows = []
#             for col in table.get("columns", []):
#                 rows.append([
#                     col.get("name"),
#                     col.get("type"),
#                     col.get("nullable"),
#                     col.get("description"),
#                 ])

#             add_table(
#                 doc,
#                 headers=["Column", "Type", "Nullable", "Description"],
#                 rows=rows
#             )

#     # =============================
#     # Cost Estimation (TABLE)
#     # =============================
#     ce = safe_get(phase3, ["cost_and_resource_estimation"], {})
#     if ce:
#         add_heading("Cost Estimation", 1)
#         rows = []
#         for c in ce.get("cost_items", []):
#             rows.append([
#                 c.get("name"),
#                 f"${c.get('monthly_estimate_usd')}",
#                 c.get("rationale"),
#             ])

#         add_table(
#             doc,
#             headers=["Cost Item", "Monthly Cost", "Rationale"],
#             rows=rows
#         )

#     # =============================
#     # Testing Strategy
#     # =============================
#     tq = safe_get(phase3, ["testing_and_qa_strategy"], {})
#     if tq:
#         add_heading("Testing & QA Strategy", 1)
#         for key, t in tq.items():
#             if isinstance(t, dict):
#                 add_paragraph(f"{t.get('name')} – {t.get('scope')}")

#     # =============================
#     # Appendices
#     # =============================
#     ap = safe_get(phase3, ["appendices"], {})
#     if ap:
#         add_heading("Appendices", 1)

#         glossary = ap.get("glossary", [])
#         if glossary:
#             add_paragraph("Glossary", bold=True)
#             for g in glossary:
#                 add_paragraph(f"{g.get('term')}: {g.get('definition')}")

#         refs = ap.get("references", [])
#         if refs:
#             add_paragraph("References", bold=True)
#             for r in refs:
#                 add_paragraph(r)

#         notes = ap.get("additional_notes")
#         if notes:
#             add_paragraph(notes)

#     doc.save(output_path)












# # =========================
# # Phase 4
# # =========================

# def ask_for_feedback() -> str | None:
#     print("\n======================================")
#     print(" Review Generated System Design ")
#     print("======================================")
#     print(
#         "Please review the generated Word document.\n"
#         "Enter feedback to refine the design.\n"
#         "Press ENTER without typing anything to finalize."
#     )
#     feedback = input("> ").strip()
#     return feedback if feedback else None